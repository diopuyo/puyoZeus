"""社長向けAIサブスクリプション福利厚生提案書を生成する。"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "AI_SUBSCRIPTION_BENEFIT_PROPOSAL_FOR_CEO_2026-08-25.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "183B56"
PALE_BLUE = "EAF3F8"
PALE_GREEN = "E8F3EC"
PALE_AMBER = "FFF4D6"
PALE_RED = "FBE9E7"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "D9E1E8"
DARK_GRAY = "4B5563"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    """セル背景色を設定する。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    """セル余白をDXAで設定する。"""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color: str = MID_GRAY, size: str = "6") -> None:
    """セルに単線罫線を設定する。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    """表見出しを改ページ後も繰り返す。"""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_no_row_split(row) -> None:
    """表の行がページをまたがないようにする。"""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    """欧文・日本語フォントを統一する。"""
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK JP")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    """narrative_proposalプリセットを適用する。"""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK JP")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for name, size, color, before, after in (
        ("Title", 26, NAVY, 0, 14),
        ("Subtitle", 13, DARK_GRAY, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK JP")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section_layout(section) -> None:
    """Letter縦・1インチ余白でページを設定する。"""
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_page_number(paragraph) -> None:
    """PAGEフィールドを追加する。"""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run_font(run, 9, color=DARK_GRAY)


def populate_header(header) -> None:
    """1種類のヘッダーを構成する。"""
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("AIサブスクリプション福利厚生　小規模実証提案")
    set_run_font(run, 9, True, NAVY)
    meta = header.add_table(rows=1, cols=2, width=Inches(6.5))
    meta.columns[0].width = Inches(3.25)
    meta.columns[1].width = Inches(3.25)
    left = meta.cell(0, 0).paragraphs[0]
    right = meta.cell(0, 1).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(left.add_run("提出先：代表取締役社長"), 8, color=DARK_GRAY)
    set_run_font(right.add_run("作成日：2026年8月25日"), 8, color=DARK_GRAY)


def populate_footer(footer_container) -> None:
    """1種類のフッターを構成する。"""
    footer = footer_container.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("社内検討用　｜　"), 8, color=DARK_GRAY)
    add_page_number(footer)


def add_header_footer(doc: Document) -> None:
    """proposal_centerpiece型のヘッダーとフッターを追加する。"""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    populate_header(section.first_page_header)
    populate_footer(section.first_page_footer)


def add_chapter_section(doc: Document) -> None:
    """章ごとの新規セクションを追加し、先頭ヘッダーを継承する。"""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section_layout(section)
    section.different_first_page_header_footer = True
    section.first_page_header.is_linked_to_previous = True
    section.first_page_footer.is_linked_to_previous = True


def add_text(doc: Document, text: str, bold_lead: str | None = None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """本文段落を追加する。"""
    p = doc.add_paragraph()
    p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, 11, True, NAVY)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, 11)
    else:
        run = p.add_run(text)
        set_run_font(run, 11)
    return p


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    """プリセット準拠の箇条書きを追加する。"""
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    set_run_font(p.add_run(text), 10.5)


def add_numbered(doc: Document, number: int, text: str) -> None:
    """番号付き段落を追加する。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(f"{number}. {text}"), 10.5)


def add_source_item(doc: Document, text: str) -> None:
    """参考資料を省スペースで追加する。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    set_run_font(p.add_run(f"• {text}"), 8.5, color=DARK_GRAY)


def add_callout(doc: Document, title: str, body: str, fill: str = PALE_BLUE, accent: str = BLUE) -> None:
    """意思決定用の強調枠を追加する。"""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 150, 190, 150, 190)
    set_cell_border(cell, accent, "10")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(title), 11.5, True, accent)
    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(0)
    set_run_font(body_p.add_run(body), 10.5, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    """比較・数値用の表を追加する。"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_ALIGN_VERTICAL.CENTER
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), 9.5, True, NAVY)
    for row_data in rows:
        row = table.add_row()
        set_no_row_split(row)
        for index, text in enumerate(row_data):
            cell = row.cells[index]
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(text), 9.3)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_metric_cards(doc: Document, cards: list[tuple[str, str]]) -> None:
    """2列の主要実績カードを追加する。"""
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    for index, (value, label) in enumerate(cards):
        cell = table.cell(index // 2, index % 2)
        set_cell_shading(cell, LIGHT_GRAY if index % 2 == 0 else PALE_BLUE)
        set_cell_margins(cell, 170, 170, 170, 170)
        set_cell_border(cell, WHITE, "12")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(value), 18, True, BLUE)
        q = cell.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.space_after = Pt(0)
        set_run_font(q.add_run(label), 9.5, color=DARK_GRAY)


def page_break(doc: Document) -> None:
    """明示改ページする。"""
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_chapter_heading(doc: Document, text: str) -> None:
    """章見出しを必ず新しいページから開始する。"""
    doc.add_heading(text, level=1)


def add_cover(doc: Document) -> None:
    """表紙を作る。"""
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("AIサブスクリプションを\n「福利厚生兼・生産性基盤」として\n導入する提案"), 24, True, NAVY)
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(sub.add_run("個人開発事例：ぷよぷよeスポーツ有利不利判定システムから得た実証"), 13, color=DARK_GRAY)
    doc.add_paragraph()
    add_callout(
        doc,
        "ご決裁いただきたい事項",
        "5名・3か月・サブスクリプション費用上限15万円の実証導入。効果とリスクを数値で確認し、継続可否は90日後に再決裁します。",
        PALE_GREEN,
        "2D7D46",
    )
    add_metric_cards(
        doc,
        [
            ("5,914件", "全体テスト基準／失敗0件（8月25日07:00時点）"),
            ("148動画", "上級者対戦データの収集完了"),
            ("99.54%", "旧基準でのSTABLE確定盤面認識実測値"),
            ("400万〜1,200万円", "同等技術資産の保守的な再構築コスト試算"),
        ],
    )
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(10)
    set_run_font(note.add_run("※ 本提案はツールの全面導入ではなく、情報管理ルール付きの小規模実証を求めるものです。"), 9, color=DARK_GRAY)


def add_executive_summary(doc: Document) -> None:
    """経営要約ページを作る。"""
    add_chapter_heading(doc, "1. 経営要約")
    add_callout(
        doc,
        "結論",
        "AIサブスクリプションは、月額利用料以上の時間を毎月わずか1〜2時間削減できれば損益分岐を超えます。福利厚生として提供することで、学習機会・採用力・定着にも波及し、単なる経費ではなく人材投資として扱えます。",
    )
    doc.add_heading("この個人開発事例が示したこと", level=2)
    add_text(doc, "本プロジェクトは、対戦動画から盤面を認識し、複数の戦略指標と機械学習で有利不利を推定し、将来は配信オーバーレイへつなぐ研究開発です。画像認識、データ収集、統計・機械学習、品質保証、長時間処理、複数AIエージェントの役割分担を一人で統合しました。")
    add_bullet(doc, "AIは、設計案の比較、実装、テスト生成、ログ解析、文書化、独立レビューの速度を高めた。")
    add_bullet(doc, "一方で、AIの出力を無条件に採用せず、既定OFF、実データ検証、独立レビュー、本番ゲートで制御した。")
    add_bullet(doc, "既存テストを通過した後も新たな重大不具合を再現し、修正前の失敗テストを追加してから是正した。これはAI導入に必要な“人間の統制”を実例で示す。")
    doc.add_heading("会社への提案", level=2)
    add_table(
        doc,
        ["項目", "提案内容", "経営上の狙い"],
        [
            ["対象", "希望者5名（職種横断）", "効果が出やすい業務を比較"],
            ["期間", "3か月", "短期で継続判断"],
            ["予算", "利用料上限15万円", "損失を限定"],
            ["合格基準", "月2時間／人以上の確認済み削減", "上限予算でも損益分岐"],
            ["統制", "法人契約・利用規程・人間レビュー", "機密・品質リスクを抑制"],
        ],
        [1.05, 2.5, 2.95],
    )


def add_case_study(doc: Document) -> None:
    """個人開発事例ページを作る。"""
    add_chapter_heading(doc, "2. 個人開発事例の概要と到達点")
    doc.add_heading("テーマ", level=2)
    add_text(doc, "上級者の『ぷよぷよeスポーツ』対戦動画を対象に、映像から6列×13行の盤面を読み取り、戦略的な有利不利を−100〜＋100で推定するシステムです。最終目標はOBS配信で使えるリアルタイム表示です。")
    doc.add_heading("技術的な広がり", level=2)
    add_table(
        doc,
        ["領域", "取り組み", "会社業務への転用可能性"],
        [
            ["画像認識", "OpenCV・CNN・時系列状態管理", "検品、帳票、映像監視、作業判定"],
            ["データ基盤", "動画収集、品質フィルタ、再生成、削除運用", "ログ分析、学習データ整備"],
            ["予測・統計", "複数指標、機械学習、A/B比較", "需要予測、リスク判定、業務改善"],
            ["品質保証", "5,900件超のテスト、実データゲート", "AI生成物の検収・監査"],
            ["開発運用", "AI間の引継ぎ、判断ログ、既定OFF", "属人化防止、安全な試行"],
        ],
        [1.05, 2.45, 3.0],
    )
    doc.add_heading("確認できている成果", level=2)
    add_bullet(doc, "STABLE確定盤面の認識は旧測定で99.54%を達成。現在構成の変更後に再測定を予定しており、旧値を無条件に流用していない。")
    add_bullet(doc, "上級者対戦148動画を収集し、未確認ティアの動画を学習から除外する品質ルールを運用。")
    add_bullet(doc, "2026年8月25日07:00時点の全体テスト基準は5,914件成功・13件スキップ・失敗0件。")
    add_bullet(doc, "独立レビューでテスト未検出の重大不具合を発見。失敗テストを先に追加し、関連150件を成功へ戻した。生成量の独立検算も115.7%から97.57%へ改善し、±5%基準を満たした。")
    add_callout(doc, "重要な留保", "製品完成とは主張しません。交換会計や表示統合には残ゲートがあり、新機能は既定OFFのままです。これは弱点ではなく、品質を守るために未完成を明示できる開発統制の証拠です。", PALE_AMBER, "B7791F")


def add_valuation(doc: Document) -> None:
    """価値評価ページを作る。"""
    add_chapter_heading(doc, "3. 本プロジェクトの価値評価")
    add_callout(doc, "中心評価", "現時点の技術資産としての再構築コストは、保守的に400万〜1,200万円相当と見積もります。これは売却価格でも将来売上でもなく、同等の検証済み資産を作り直す費用の目安です。")
    doc.add_heading("見積もり方法", level=2)
    add_table(
        doc,
        ["前提", "保守ケース", "上限ケース"],
        [
            ["同等資産の開発工数", "4人月", "8人月"],
            ["技術者の完全原価", "90万円／人月", "150万円／人月"],
            ["計算値", "360万円", "1,200万円"],
            ["提示レンジ", "約400万円", "約1,200万円"],
        ],
        [2.2, 2.15, 2.15],
    )
    add_text(doc, "工数には、画像認識、時系列処理、学習データ整備、機械学習、テスト、実データ検証、開発記録の再構築を含みます。実際の投下時間や給与額を積み上げた原価ではなく、第三者が同等水準へ到達する場合の置換費用です。")
    doc.add_heading("価値を4つに分解", level=2)
    add_table(
        doc,
        ["価値の種類", "現時点の評価", "根拠／留保"],
        [
            ["技術資産価値", "高い", "複数技術領域、実データ、テスト、運用記録が一体"],
            ["ポートフォリオ価値", "非常に高い", "課題発見から検証・失敗・是正まで説明可能"],
            ["採用・広報価値", "中〜高", "AI活用と品質統制を具体例で発信できる"],
            ["製品・売上価値", "未確定", "本番統合・権利整理・利用者検証が未完了"],
        ],
        [1.45, 1.35, 3.7],
    )


def add_ai_effect(doc: Document) -> None:
    """AI活用の意味とブログ方針を説明する。"""
    add_chapter_heading(doc, "4. AIサブスクリプションが生んだ価値")
    doc.add_heading("速さだけでなく、思考の幅と品質確認を増やした", level=2)
    add_table(
        doc,
        ["AIの役割", "具体的な使い方", "得られた価値"],
        [
            ["設計者", "方式比較、境界条件、失敗仮説の列挙", "検討漏れを減らす"],
            ["実装者", "定型コード、テスト、診断スクリプト", "反復速度を高める"],
            ["分析者", "ログ集計、異常分解、A/B整理", "原因特定を早める"],
            ["レビュー補助", "別AIによる独立確認と再現", "思い込みと迎合を抑える"],
            ["記録者", "引継ぎ文書、判断ログ、ロードマップ", "継続性を保つ"],
        ],
        [1.2, 2.5, 2.8],
    )
    add_callout(doc, "実例から得た原則", "AIは“正解を出す装置”ではなく、“試行回数と検証観点を増やす同僚”として使うと価値が出ます。採用判断は人間が持ち、テスト・実データ・権限管理で囲う必要があります。", PALE_GREEN, "2D7D46")
    doc.add_heading("ブログ／ポートフォリオで伝えるべき構成", level=2)
    add_numbered(doc, 1, "問題設定：なぜ動画から有利不利を測るのが難しいか。")
    add_numbered(doc, 2, "設計：画像認識、時系列状態、指標、機械学習をどう分離したか。")
    add_numbered(doc, 3, "AI協働：AIごとの役割、引継ぎ、独立レビューをどう運用したか。")
    add_numbered(doc, 4, "失敗：99%から1%へ急落する現象や、既存テスト後に見つかった欠陥をどう再現したか。")
    add_numbered(doc, 5, "品質統制：既定OFF、実データゲート、採用フラグの単一情報源。")
    add_numbered(doc, 6, "成果と未完成：数値成果と残課題を同時に示す。")
    add_text(doc, "公開時は、第三者動画の映像・ユーザー名・URL・ローカルパス・秘密情報を除き、必要なら自作図と集計値へ置き換えます。ゲーム映像の転載可否は公開先の規約と権利条件を別途確認します。")


def add_roi(doc: Document) -> None:
    """ROI試算ページを作る。"""
    add_chapter_heading(doc, "5. 福利厚生としての投資対効果")
    add_text(doc, "AIサブスクリプションを『自由に使える福利厚生』だけで終わらせず、本人の学習と会社の生産性向上を同時に測る制度にします。以下は保証値ではなく、意思決定のための仮定です。")
    doc.add_heading("標準シナリオ（5名・年間換算）", level=2)
    add_table(
        doc,
        ["1人あたり削減", "年間便益", "年間利用料", "年間純便益", "ROI"],
        [
            ["2時間／月", "60万円", "36万円", "24万円", "67%"],
            ["4時間／月", "120万円", "36万円", "84万円", "233%"],
            ["8時間／月", "240万円", "36万円", "204万円", "567%"],
        ],
        [1.3, 1.3, 1.3, 1.35, 1.25],
    )
    add_text(doc, "計算前提：5名、完全原価5,000円／時間、利用料6,000円／人・月。ROI＝（便益−利用料）÷利用料。品質悪化、手戻り、管理工数があれば便益から差し引きます。")
    add_callout(doc, "損益分岐", "標準前提では1.2時間／人・月の削減で利用料を回収します。実証予算の上限15万円を全額使った場合でも、2時間／人・月で3か月の便益15万円となり損益分岐です。", PALE_GREEN, "2D7D46")
    doc.add_heading("参考となる法人向け価格", level=2)
    add_bullet(doc, "ChatGPT Business：年払い20米ドル／ユーザー・月、月払い25米ドル／ユーザー・月（公式価格ページ、2026年8月25日確認）。")
    add_bullet(doc, "GitHub Copilot Business：19米ドル／付与席・月（公式ドキュメント、2026年8月25日確認）。")
    add_text(doc, "為替、税、最低席数、契約条件、必要機能が変わるため、本提案では個別製品を決め打ちせず、円建て上限予算で管理します。")
    doc.add_heading("福利厚生としての追加便益", level=2)
    add_bullet(doc, "社員が個人負担なしで最新スキルを試せるため、学習格差と“隠れAI利用”を減らす。")
    add_bullet(doc, "採用候補者へ、学習投資と新技術の安全運用を重視する会社であることを示せる。")
    add_bullet(doc, "日常業務で得たプロンプト、手順、注意点を社内共有資産へ変換できる。")


def add_pilot(doc: Document) -> None:
    """実証計画ページを作る。"""
    add_chapter_heading(doc, "6. 3か月実証の設計")
    add_callout(doc, "提案条件", "対象5名、期間3か月、サブスクリプション費用上限15万円。法人向けプランを使い、機密情報の扱いと人間レビューを事前に定めます。")
    doc.add_heading("実施ステップ", level=2)
    add_numbered(doc, 1, "0週目：対象業務、基準時間、品質指標、禁止情報を登録する。")
    add_numbered(doc, 2, "1週目：90分の利用研修を行い、良い依頼方法、誤りの確認、引用・権利、機密情報を共有する。")
    add_numbered(doc, 3, "2〜11週目：週次で利用例、削減時間、手戻り、品質変化を簡潔に記録する。")
    add_numbered(doc, 4, "6週目：中間確認。利用されない席は対象者を入れ替える。重大リスクがあれば停止する。")
    add_numbered(doc, 5, "12週目：効果、事故、利用率、継続候補業務をまとめ、社長へ継続・縮小・終了を提案する。")
    doc.add_heading("測定指標と合格条件", level=2)
    add_table(
        doc,
        ["指標", "測定方法", "合格条件"],
        [
            ["確認済み削減時間", "導入前後の実作業時間＋上長確認", "中央値2時間／人・月以上"],
            ["品質", "手戻り率、誤送信、レビュー指摘", "導入前より悪化しない"],
            ["利用率", "月1回以上の業務利用", "対象者の60%以上"],
            ["安全性", "機密・個人情報・権利事故", "重大事故0件"],
            ["再利用資産", "共有テンプレート・手順", "有用例を5件以上蓄積"],
        ],
        [1.35, 3.0, 2.15],
    )
    add_text(doc, "“AIを使った回数”ではなく、確認できた時間、品質、再利用資産で評価します。成果が出なければ3か月で停止できるため、導入リスクは上限15万円に限定されます。")


def add_governance(doc: Document) -> None:
    """リスク管理ページを作る。"""
    add_chapter_heading(doc, "7. リスクと統制")
    add_table(
        doc,
        ["主なリスク", "想定される問題", "実証時の統制"],
        [
            ["機密・個人情報", "外部サービスへ不適切に入力", "法人契約、入力禁止区分、匿名化、管理者設定"],
            ["誤情報", "もっともらしい誤答を採用", "人間レビュー、根拠確認、重要判断への単独利用禁止"],
            ["著作権・知財", "生成物や入力物の権利が不明", "出典記録、転載禁止、公開前レビュー"],
            ["シャドーAI", "個人アカウントで無統制利用", "承認済み環境を提供し、相談窓口を設置"],
            ["費用膨張", "席や追加利用の増加", "上限予算、月次棚卸し、未利用席停止"],
            ["依存・技能低下", "検証せずAIへ丸投げ", "成果物責任は本人、基礎技能とレビューを評価"],
        ],
        [1.3, 2.45, 2.75],
    )
    doc.add_heading("本プロジェクトで実践した統制", level=2)
    add_bullet(doc, "新機能は既定OFF。実データの合格条件を満たすまで本番表示へ接続しない。")
    add_bullet(doc, "採用フラグを一か所で管理し、採用日と根拠を記録する。")
    add_bullet(doc, "別AIによる独立レビューで、既存テストが見逃した重大不具合を最小再現する。")
    add_bullet(doc, "修正前に失敗テストを追加し、改善後も全体回帰を行う。")
    add_bullet(doc, "未完成・未測定・旧測定値を区別し、都合のよい数字だけを出さない。")
    add_callout(doc, "経営上の意味", "AIを禁止しても利用は地下化しやすく、統制も学習も残りません。承認済み環境と小さな予算を用意し、可視化された実証として管理する方が安全です。", PALE_AMBER, "B7791F")


def add_decision(doc: Document) -> None:
    """決裁依頼と参考情報をまとめる。"""
    add_chapter_heading(doc, "8. 決裁依頼と次の一手")
    add_callout(
        doc,
        "決裁依頼",
        "AIサブスクリプション福利厚生の3か月実証を、対象5名・利用料上限15万円・法人向け契約・利用規程付きで承認いただきたい。90日後に数値報告を行い、継続は別途決裁とします。",
        PALE_GREEN,
        "2D7D46",
    )
    doc.add_heading("承認後2週間で用意するもの", level=2)
    add_bullet(doc, "対象者と対象業務の一覧")
    add_bullet(doc, "利用規程1枚、機密区分、公開前レビュー手順")
    add_bullet(doc, "導入前の作業時間・品質基準")
    add_bullet(doc, "週次記録テンプレートと90日報告書のひな型")
    doc.add_heading("経営判断のポイント", level=2)
    add_text(doc, "本提案の本質は特定製品の購入ではなく、社員がAIを安全に試し、成果を数字と再利用資産へ変える仕組みへの小規模投資です。本事例は、複雑な課題でも効果が見込めることと、品質統制が不可欠であることの双方を示します。")
    add_text(doc, "上限15万円で3か月後に撤退でき、標準ケースでは月1.2時間／人、上限予算でも月2時間／人の削減で損益分岐です。効果が確認できた業務だけへ広げます。")
    doc.add_heading("参考資料・前提", level=2)
    add_source_item(doc, "社内根拠：AGENTS.md、PROJECT_STATE.md、agent_coordination配下の進捗・レビュー記録（2026年8月25日確認）")
    add_source_item(doc, "OpenAI Business価格：openai.com/business/pricing（2026年8月25日確認）")
    add_source_item(doc, "GitHub Copilotプラン：docs.github.com/en/copilot/get-started/plans（同日確認）")
    add_source_item(doc, "価値試算：4〜8人月×90万〜150万円／人月。実売価値・将来売上は含めない")
    add_source_item(doc, "ROI試算：完全原価5,000円／時、利用料6,000円／人・月。実証時に実額へ置換")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("価格・契約条件は変更され得るため、契約前に見積書、利用規約、データ取扱条件を再確認します。"), 8.5, color=DARK_GRAY)


def build_document() -> Path:
    """提案書を組み立てて保存する。"""
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = False
    configure_styles(doc)
    configure_section_layout(doc.sections[0])
    add_header_footer(doc)
    props = doc.core_properties
    props.title = "AIサブスクリプション福利厚生の小規模実証提案"
    props.subject = "個人開発事例に基づくAI人材投資・ROI提案"
    props.author = "個人開発者"
    props.keywords = "AI, 福利厚生, 生産性, ポートフォリオ, ROI, 小規模実証"
    add_cover(doc)
    for builder in (
        add_executive_summary,
        add_case_study,
        add_valuation,
        add_ai_effect,
        add_roi,
        add_pilot,
        add_governance,
        add_decision,
    ):
        add_chapter_section(doc)
        builder(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
